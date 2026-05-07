"""
Scientific Keyphrase Extractor — Advanced (PageIndex Edition)

Key change: replaces TF-IDF keyword scoring with PageIndex hierarchical
tree indexing + LLM-reasoning-based keyphrase extraction, powered by
Google Gemini via the GEMINI_API_KEY Streamlit secret.

Users see no API key input — authentication, extraction, visualisations,
and activity history all work out of the box once deployed.
"""

import re
import sys
import os
import json
import tempfile
import hashlib
import io
import base64
from datetime import datetime
from pathlib import Path

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import streamlit as st
from wordcloud import WordCloud

# Document parsers (unchanged from original)
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Auth / DB (unchanged)
import sqlite3
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC

# ── PageIndex integration ──────────────────────────────────────────────────────
# Add the project root to sys.path so the local pageindex package is importable
sys.path.insert(0, os.path.dirname(__file__))
from pageindex import PageIndexClient
from pageindex.utils import llm_completion

# ── Gemini API key — loaded once from Streamlit Secrets ───────────────────────
# Stored under Settings → Secrets in Streamlit Cloud as:
#   GEMINI_API_KEY = "AIza..."
# LiteLLM (used internally by PageIndex) picks up GEMINI_API_KEY automatically
# for any model string prefixed with "gemini/".

def _load_gemini_key() -> str:
    """
    Read the Gemini API key from st.secrets (Streamlit Cloud deployment)
    with a graceful fallback to the local environment variable for local dev.
    Raises a clear error if neither source provides a key.
    """
    key = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY", "")
    if not key:
        st.error(
            "⚠️ **GEMINI_API_KEY is not configured.**  "
            "If you are the app owner, please add it under "
            "*Settings → Secrets* in Streamlit Cloud."
        )
        st.stop()
    return key


# ── Gemini model to use throughout the app ────────────────────────────────────
GEMINI_MODEL = "gemini/gemini-2.0-flash"


# ── Security and Database Setup (unchanged) ────────────────────────────────────

def setup_database():
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (username TEXT PRIMARY KEY, password_hash TEXT, salt TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS activities
                 (username TEXT, timestamp TEXT, file_name TEXT, keyphrases TEXT,
                  FOREIGN KEY (username) REFERENCES users(username))''')
    conn.commit()
    return conn

def generate_key(password: str, salt: bytes) -> bytes:
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=100000,
    )
    return base64.urlsafe_b64encode(kdf.derive(password.encode()))

def encrypt_data(data: str, key: bytes) -> str:
    f = Fernet(key)
    return f.encrypt(data.encode()).decode()

def decrypt_data(encrypted_data: str, key: bytes) -> str:
    f = Fernet(key)
    return f.decrypt(encrypted_data.encode()).decode()

def register_user(username: str, password: str, conn):
    salt = os.urandom(16)
    key = generate_key(password, salt)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO users (username, password_hash, salt) VALUES (?, ?, ?)",
                 (username, key.decode(), salt.hex()))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False

def verify_user(username: str, password: str, conn):
    c = conn.cursor()
    c.execute("SELECT password_hash, salt FROM users WHERE username=?", (username,))
    result = c.fetchone()
    if result:
        stored_key = result[0].encode()
        salt = bytes.fromhex(result[1])
        key = generate_key(password, salt)
        return key.decode() == stored_key.decode()
    return False

def save_activity(username: str, file_name: str, keyphrases_df: pd.DataFrame, conn):
    c = conn.cursor()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    keyphrases_json = keyphrases_df.to_json()
    c.execute(
        "INSERT INTO activities (username, timestamp, file_name, keyphrases) VALUES (?, ?, ?, ?)",
        (username, timestamp, file_name, keyphrases_json)
    )
    conn.commit()

def get_user_activities(username: str, conn):
    c = conn.cursor()
    c.execute(
        "SELECT timestamp, file_name, keyphrases FROM activities WHERE username=? ORDER BY timestamp DESC",
        (username,)
    )
    return c.fetchall()

# ── Text extraction helpers (unchanged) ────────────────────────────────────────

def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    return ''.join(page.extract_text() or '' for page in reader.pages)

def extract_text_from_docx(docx_path):
    doc = DocxDocument(docx_path)
    return '\n'.join(para.text for para in doc.paragraphs)

def preprocess_text(text: str) -> str:
    """Light cleanup used only for frequency counting (not for indexing)."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    return text.lower()

# ── DOCX → Markdown converter for PageIndex MD mode ───────────────────────────

def docx_to_markdown(docx_path: str, out_md_path: str):
    """
    Convert a DOCX file to a minimal Markdown file so PageIndex can index it
    using its heading-aware MD mode.  Heading styles (Heading 1–6) are mapped
    to the corresponding Markdown '#' prefix; all other paragraphs are output
    as plain text.
    """
    doc = DocxDocument(docx_path)
    heading_map = {
        'Heading 1': '#', 'Heading 2': '##', 'Heading 3': '###',
        'Heading 4': '####', 'Heading 5': '#####', 'Heading 6': '######',
    }
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        prefix = heading_map.get(para.style.name, '')
        lines.append(f"{prefix} {text}" if prefix else text)
    with open(out_md_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(lines))

# ── PageIndex-based keyphrase extraction ───────────────────────────────────────
# This is the ONLY section that replaces the original compute_tfidf() logic.

def _flatten_tree(nodes: list, depth: int = 0) -> list[dict]:
    """Recursively collect every node's title and summary from the PageIndex tree."""
    flat = []
    for node in nodes:
        title   = (node.get('title') or '').strip()
        summary = (node.get('summary') or '').strip()
        if title:
            flat.append({'title': title, 'summary': summary, 'depth': depth})
        if node.get('nodes'):
            flat.extend(_flatten_tree(node['nodes'], depth + 1))
    return flat


def extract_keyphrases_pageindex(
    file_path: str,
    file_type: str,
    model: str,
    num_keyphrases: int,
    original_text: str,
) -> pd.DataFrame:
    """
    Replace TF-IDF with PageIndex reasoning-based extraction.

    The GEMINI_API_KEY environment variable must already be set before this
    function is called (done once at app startup via _load_gemini_key()).
    LiteLLM routes any "gemini/" prefixed model string through that key
    automatically — no OpenAI key is required.

    Steps
    -----
    1. Index the document with PageIndexClient → builds a hierarchical
       tree (section titles + LLM-generated summaries).
    2. Flatten the tree to gather all section context.
    3. Send that context to the LLM asking it to extract and score the
       top N scientific keyphrases with reasoning.
    4. Count each keyphrase's occurrence in the original text to fill
       the 'Frequency' column (preserving compatibility with existing
       visualisations).
    5. Return a DataFrame with columns: KeyPhrase, Frequency, Importance (%).
    """
    # ── Step 1: Index the document ────────────────────────────────────────
    # Do NOT pass api_key to PageIndexClient — it would set OPENAI_API_KEY,
    # which is irrelevant here. Instead the GEMINI_API_KEY env var (already
    # set at app startup) is picked up transparently by LiteLLM.
    client = PageIndexClient(model=model)

    if file_type == "pdf":
        doc_id = client.index(file_path, mode="pdf")
    else:  # DOCX / DOC → convert to MD first
        md_temp = file_path.replace('.docx', '.md').replace('.doc', '.md')
        docx_to_markdown(file_path, md_temp)
        doc_id = client.index(md_temp, mode="md")

    # ── Step 2: Retrieve and flatten the tree ────────────────────────────
    structure_json = client.get_document_structure(doc_id)
    structure      = json.loads(structure_json)
    tree_nodes     = _flatten_tree(structure)

    if not tree_nodes:
        raise ValueError(
            "PageIndex returned an empty tree.  "
            "The document may be too short, image-only, or failed to parse."
        )

    # Build a compact textual representation of the tree for the LLM prompt
    tree_text = "\n".join(
        f"[L{n['depth']}] {n['title']}"
        + (f" — {n['summary'][:200]}" if n['summary'] else "")
        for n in tree_nodes
    )

    # ── Step 3: LLM-based keyphrase scoring ──────────────────────────────
    prompt = f"""You are an expert scientific document analyst.
Below is the hierarchical structure (section titles and summaries) of a scientific document.

DOCUMENT STRUCTURE:
{tree_text}

Your task:
Extract exactly {num_keyphrases} important scientific keyphrases from this document structure.
For each keyphrase, provide an importance score from 1–100 reflecting how central
and distinctive it is to the scientific content of this document.

Rules:
- Prefer multi-word phrases (1–4 words) that are domain-specific and meaningful.
- Do NOT include generic words like "introduction", "conclusion", "results", "figure".
- Each keyphrase must be unique.

Return ONLY a valid JSON array, no preamble, no markdown fences. Example format:
[
  {{"keyphrase": "neural architecture search", "importance": 95}},
  {{"keyphrase": "gradient descent", "importance": 82}}
]
"""
    response = llm_completion(model=model, prompt=prompt)

    # Parse LLM response robustly
    try:
        cleaned = re.sub(r'```(?:json)?|```', '', response).strip()
        keyphrases_raw = json.loads(cleaned)
    except json.JSONDecodeError:
        # Fallback: extract any JSON array from the response
        match = re.search(r'\[.*\]', response, re.DOTALL)
        if match:
            keyphrases_raw = json.loads(match.group())
        else:
            raise ValueError(f"LLM returned unparseable output:\n{response}")

    # ── Step 4: Build the DataFrame ───────────────────────────────────────
    lower_text = preprocess_text(original_text)
    rows = []
    for item in keyphrases_raw[:num_keyphrases]:
        phrase     = str(item.get("keyphrase", "")).strip().lower()
        importance = float(item.get("importance", 0))
        if not phrase:
            continue
        # Count occurrences in the cleaned original text
        frequency = len(re.findall(re.escape(phrase), lower_text))
        rows.append({"KeyPhrase": phrase, "Frequency": frequency, "_raw_importance": importance})

    if not rows:
        raise ValueError("No keyphrases could be extracted from the LLM response.")

    df = pd.DataFrame(rows)

    # ── Step 5: Normalise importance to percentage ────────────────────────
    max_imp = df["_raw_importance"].max() or 1.0
    df["Importance (%)"] = (df["_raw_importance"] / max_imp * 100).round(2)
    df.drop(columns=["_raw_importance"], inplace=True)
    df.sort_values("Importance (%)", ascending=False, inplace=True)
    df.reset_index(drop=True, inplace=True)

    return df[["KeyPhrase", "Frequency", "Importance (%)"]]


# ── Shared visualisation helper (unchanged logic, extracted to avoid duplication) ──

def render_visualisations(df: pd.DataFrame):
    """Render all four charts for a keyphrases DataFrame."""
    # Word Cloud
    st.subheader("Keyphrase Word Cloud")
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(
        df.set_index('KeyPhrase')['Importance (%)'].to_dict()
    )
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.axis("off")
    st.pyplot(fig)
    plt.close(fig)

    # Donut Pie Chart
    st.subheader("Keyphrase Importance (Donut Chart)")
    fig1, ax1 = plt.subplots()
    ax1.pie(df['Importance (%)'], labels=df['KeyPhrase'], autopct='%1.1f%%', startangle=140)
    ax1.axis('equal')
    centre_circle = plt.Circle((0, 0), 0.70, fc='white')
    fig1.gca().add_artist(centre_circle)
    st.pyplot(fig1)
    plt.close(fig1)

    # Radial Bar Chart
    st.subheader("Radial Keyphrase Bar Chart")
    angles = np.linspace(0, 2 * np.pi, len(df), endpoint=False)
    fig2, ax2 = plt.subplots(figsize=(8, 8), subplot_kw={'polar': True})
    ax2.bar(angles, df['Frequency'],
            color=plt.cm.viridis(df['Importance (%)'] / 100), alpha=0.7)
    ax2.set_xticks(angles)
    ax2.set_xticklabels(df['KeyPhrase'], fontsize=8, rotation=45)
    st.pyplot(fig2)
    plt.close(fig2)

    # Frequency Histogram
    st.subheader("Keyphrase Frequency (Histogram)")
    fig3, ax3 = plt.subplots(figsize=(10, 5))
    sns.barplot(x=df['KeyPhrase'], y=df['Frequency'], palette='viridis', ax=ax3)
    ax3.set_xticklabels(ax3.get_xticklabels(), rotation=45, ha="right")
    ax3.set_title("Keyphrase Frequency")
    ax3.set_ylabel("Frequency")
    ax3.set_xlabel("KeyPhrase")
    st.pyplot(fig3)
    plt.close(fig3)


# ── Main Streamlit application ────────────────────────────────────────────────

def main():
    st.title("Scientific Keyphrase Extractor")
    st.caption("Powered by **Google Gemini + PageIndex** — hierarchical tree indexing + LLM-reasoning-based retrieval")

    # Load Gemini API key from Streamlit Secrets and inject into the environment
    # so LiteLLM (used by PageIndex) can route gemini/ model strings correctly.
    gemini_key = _load_gemini_key()
    os.environ["GEMINI_API_KEY"] = gemini_key

    conn = setup_database()

    if 'username' not in st.session_state:
        st.session_state.username = None

    # ── Login / Register (unchanged) ─────────────────────────────────────
    if not st.session_state.username:
        st.subheader("Login / Register")
        tab1, tab2 = st.tabs(["Login", "Register"])

        with tab1:
            login_username = st.text_input("Username")
            login_password = st.text_input("Password", type="password")
            if st.button("Login"):
                if verify_user(login_username, login_password, conn):
                    st.session_state.username = login_username
                    st.success("Login successful!")
                else:
                    st.error("Invalid username or password")

        with tab2:
            reg_username = st.text_input("New Username")
            reg_password = st.text_input("New Password", type="password")
            reg_confirm  = st.text_input("Confirm Password", type="password")
            if st.button("Register"):
                if reg_password != reg_confirm:
                    st.error("Passwords do not match!")
                elif len(reg_password) < 6:
                    st.error("Password must be at least 6 characters long!")
                else:
                    if register_user(reg_username, reg_password, conn):
                        st.success("Registration successful! Please login.")
                    else:
                        st.error("Username already exists!")
        return

    # ── Authenticated view ────────────────────────────────────────────────
    st.write(f"Welcome, **{st.session_state.username}**!")
    if st.button("Logout"):
        st.session_state.username = None
        st.rerun()

    st.write("Extract keyphrases from scientific documents with stunning visuals!")

    # ── Sidebar: informational only — no key input needed ────────────────
    with st.sidebar:
        st.header("⚙️ About This App")
        st.info(
            "**Powered by Google Gemini + PageIndex**\n\n"
            "1. Your document is parsed into a hierarchical *Table-of-Contents* "
            "tree where each node carries a section title and a Gemini-generated summary.\n"
            "2. Gemini then reasons over the full tree to extract and rank the most "
            "scientifically significant keyphrases — no chunking, no vector DB.\n\n"
            f"🤖 Model in use: `{GEMINI_MODEL}`"
        )
        st.markdown("---")
        st.caption(
            "The Gemini API key is pre-configured by the app owner and is never "
            "visible to users."
        )

    # ── File upload and extraction ────────────────────────────────────────
    uploaded_file = st.file_uploader(
        "Drag / Drop / Upload Your Scientific Content (PDF / Doc / Docx)",
        type=["pdf", "docx", "doc"],
    )
    num_keyphrases = st.slider(
        "Select Number of Keyphrases to Extract", min_value=5, max_value=50, value=20, step=5
    )

    if st.button("Extract Keyphrases"):
        if not uploaded_file:
            st.error("Please upload a valid file!")
            return

        file_type = uploaded_file.name.rsplit('.', 1)[-1].lower()
        if file_type not in ("pdf", "doc", "docx"):
            st.error("Unsupported file format!")
            return

        with st.spinner("Building PageIndex tree and extracting keyphrases with Gemini reasoning…"):
            # Write the uploaded file to a temp path so PageIndex can read it
            suffix = f".{file_type}"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name

            try:
                # Extract raw text for frequency counting (unchanged helper)
                if file_type == "pdf":
                    original_text = extract_text_from_pdf(tmp_path)
                else:
                    original_text = extract_text_from_docx(tmp_path)

                # ── Core replacement: PageIndex instead of TF-IDF ─────────
                result_df = extract_keyphrases_pageindex(
                    file_path      = tmp_path,
                    file_type      = file_type if file_type == "pdf" else "docx",
                    model          = GEMINI_MODEL,
                    num_keyphrases = num_keyphrases,
                    original_text  = original_text,
                )
                # ─────────────────────────────────────────────────────────

            except Exception as exc:
                st.error(f"Extraction failed: {exc}")
                return
            finally:
                os.unlink(tmp_path)

        st.success("Keyphrase Extraction Successful!")

        # ── Results display (unchanged) ───────────────────────────────────
        st.subheader("Keyphrases")
        st.dataframe(result_df)

        # Save activity
        save_activity(st.session_state.username, uploaded_file.name, result_df, conn)

        render_visualisations(result_df)

    # ── Past Activities (unchanged) ───────────────────────────────────────
    st.subheader("Your Past Activities")
    activities = get_user_activities(st.session_state.username, conn)

    if activities:
        for timestamp, file_name, keyphrases_json in activities:
            with st.expander(f"{file_name} — {timestamp}"):
                keyphrases_df = pd.read_json(keyphrases_json)
                st.dataframe(keyphrases_df)
                render_visualisations(keyphrases_df)
    else:
        st.info("No past activities found.")


if __name__ == "__main__":
    main()
